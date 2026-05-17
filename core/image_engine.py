import io
import os

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from .docx_utils import decode_image_stream, get_image_dimensions_cm


class ImageEngine:
    """
    ImageEngine V2 — Pemrosesan Gambar Proposional & Stabil.
    Menggunakan rasio aspek asli agar tidak gepeng/stretch.
    """
    def __init__(self, document):
        self.document = document

    def insert_cover_image(self, image_source, max_width_cm: float = 6.0, max_height_cm: float = 6.0):
        """Sisipkan gambar cover (biasanya logo sekolah) dengan resize proporsional."""
        if not image_source:
            return None
            
        stream = decode_image_stream(image_source)
        if not stream:
            return None

        dims = get_image_dimensions_cm(stream)
        w, h = self._calculate_dimensions(dims, max_width_cm, max_height_cm)

        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(stream, width=Cm(w), height=Cm(h))
        
        return paragraph

    def insert_content_image(self, image_source, caption: str = "", max_width_cm: float = 14.0, max_height_cm: float = 10.0, toc_entry_text: str = ""):
        """Sisipkan gambar ke dalam isi laporan dengan caption."""
        if not image_source:
            return None

        stream = decode_image_stream(image_source)
        if not stream:
            return None

        dims = get_image_dimensions_cm(stream)
        w, h = self._calculate_dimensions(dims, max_width_cm, max_height_cm)

        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(stream, width=Cm(w), height=Cm(h))

        if caption:
            from .styles import apply_caption_style
            cap_para = self.document.add_paragraph()
            apply_caption_style(cap_para)
            cap_para.add_run(caption.strip())
            
        return paragraph

    def insert_image_fit(self, image_source, max_width_cm: float, max_height_cm: float, paragraph=None):
        """Sisipkan gambar (misal TTD) ke paragraf tertentu agar pas di area tersedia."""
        if not image_source:
            return None
        
        stream = decode_image_stream(image_source)
        if not stream:
            return None

        dims = get_image_dimensions_cm(stream)
        w, h = self._calculate_dimensions(dims, max_width_cm, max_height_cm)

        if paragraph is None:
            paragraph = self.document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        run = paragraph.add_run()
        run.add_picture(stream, width=Cm(w), height=Cm(h))
        return paragraph

    def _calculate_dimensions(self, original_dims, max_w, max_h):
        """Hitung width/height agar tetap proporsional sesuai limit."""
        if not original_dims:
            return max_w, max_h
            
        orig_w, orig_h = original_dims
        ratio = min(max_w / orig_w, max_h / orig_h)
        
        # Jika gambar asli lebih kecil dari limit, jangan di-stretch ke besar
        final_ratio = min(ratio, 1.0) 
        
        return orig_w * final_ratio, orig_h * final_ratio
