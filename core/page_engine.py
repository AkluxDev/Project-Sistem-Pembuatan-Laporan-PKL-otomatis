from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from .docx_utils import add_field_run, clear_paragraph
from .styles import FONT_NAME, MARGIN_BOTTOM_CM, MARGIN_LEFT_CM, MARGIN_RIGHT_CM, MARGIN_TOP_CM


class PageEngine:
    def __init__(self, document):
        self.document = document
        self._configure_section(self.document.sections[0])

    def configure_base_section(self):
        self._configure_section(self.document.sections[0])
        return self.document.sections[0]

    def add_section(self, start=WD_SECTION_START.NEW_PAGE):
        section = self.document.add_section(start)
        self._configure_section(section)
        return section

    def add_page_break(self):
        self.document.add_page_break()

    def _configure_section(self, section):
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(MARGIN_LEFT_CM)
        section.top_margin = Cm(MARGIN_TOP_CM)
        section.right_margin = Cm(MARGIN_RIGHT_CM)
        section.bottom_margin = Cm(MARGIN_BOTTOM_CM)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.5)
        section.different_first_page_header_footer = False
        return section

    def setup_cover_section(self):
        section = self.document.sections[0]
        section.different_first_page_header_footer = True
        self._clear_footer(section.footer)
        self._clear_footer(section.first_page_footer)
        self._set_page_number_type(section, start=1, fmt="decimal")
        return section

    def setup_preliminary_section(self, section, start=1):
        """Setup section untuk Kata Pengantar & Daftar Isi (i, ii, iii...)."""
        self._attach_footer(section, page_format="lowerRoman", start=start)

    def setup_body_section(self, section, start=1):
        """Setup section untuk konten utama BAB I dst (1, 2, 3...)."""
        self._attach_footer(section, page_format="decimal", start=start)

    def _attach_footer(self, section, page_format: str, start: int):
        """Sisipkan nomor halaman menggunakan field PAGE yang dinamis."""
        section.footer.is_linked_to_previous = False
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        clear_paragraph(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Injeksi field PAGE secara manual melalui XML agar stabil
        run = paragraph.add_run()
        
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_end)
        
        # Set tipe penomoran dan angka mulai pada property section
        self._set_page_number_type(section, start=start, fmt=page_format)

    def _clear_footer(self, footer):
        footer.is_linked_to_previous = False
        for paragraph in footer.paragraphs:
            clear_paragraph(paragraph)

    def _set_page_number_type(self, section, start: int, fmt: str):
        sect_pr = section._sectPr
        page_num = sect_pr.find(qn("w:pgNumType"))
        if page_num is None:
            page_num = OxmlElement("w:pgNumType")
            sect_pr.append(page_num)
        page_num.set(qn("w:start"), str(start))
        page_num.set(qn("w:fmt"), fmt)
