from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .docx_utils import sanitize_paragraph_lines
from .styles import (
    BODY_SIZE,
    HEADING_SIZE,
    SUBHEADING_SIZE,
    apply_heading_style,
    apply_list_style,
    apply_normal_style,
    apply_signature_style,
    apply_subheading_style,
    format_run,
)


def enable_flow_control(paragraph, *, keep_next=False, keep_lines=False, widow_control=True, page_break_before=False):
    p_pr = paragraph._p.get_or_add_pPr()
    _set_on_off(p_pr, "w:keepNext", keep_next)
    _set_on_off(p_pr, "w:keepLines", keep_lines)
    _set_on_off(p_pr, "w:widowControl", widow_control)
    _set_on_off(p_pr, "w:pageBreakBefore", page_break_before)
    return paragraph


def _set_on_off(parent, tag: str, enabled: bool):
    child = parent.find(qn(tag))
    if not enabled:
        if child is not None:
            parent.remove(child)
        return
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    child.set(qn("w:val"), "true")


class ParagraphEngine:
    def __init__(self, document):
        self.document = document

    def add_body_text(self, text: str):
        paragraphs = []
        for block in sanitize_paragraph_lines(text):
            paragraph = self.document.add_paragraph()
            apply_normal_style(paragraph)
            enable_flow_control(paragraph, keep_lines=True)
            if len(block) < 85:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(block)
            format_run(run, size=BODY_SIZE, italic=False, bold=False)
            paragraphs.append(paragraph)
        return paragraphs

    def add_heading_bab(self, label: str, title: str | None = None):
        paragraph = self.document.add_paragraph()
        apply_heading_style(paragraph, level=1)
        enable_flow_control(paragraph, keep_next=True, keep_lines=True)
        heading_text = label.upper() if not title else f"{label.upper()} {title.upper()}"
        format_run(paragraph.add_run(heading_text), bold=True, italic=False, size=HEADING_SIZE)
        return [paragraph]

    def add_section_heading(self, text: str):
        paragraph = self.document.add_paragraph()
        apply_heading_style(paragraph, level=1)
        enable_flow_control(paragraph, keep_next=True, keep_lines=True)
        format_run(paragraph.add_run(text.upper()), bold=True, italic=False, size=HEADING_SIZE)
        return paragraph

    def add_subheading(self, text: str, level: int = 2):
        paragraph = self.document.add_paragraph()
        apply_subheading_style(paragraph, level=level)
        enable_flow_control(paragraph, keep_next=True, keep_lines=True)
        format_run(paragraph.add_run(text), bold=True, italic=False, size=SUBHEADING_SIZE)
        return paragraph

    def add_centered_line(self, text: str, *, bold=False, size=None):
        paragraph = self.document.add_paragraph()
        apply_signature_style(paragraph)
        enable_flow_control(paragraph, keep_lines=True)
        format_run(paragraph.add_run(text), bold=bold, italic=False, size=size or BODY_SIZE)
        return paragraph

    def add_list_item(self, label: str, title: str, text: str, level: int = 0):
        paragraph = self.document.add_paragraph()
        apply_list_style(paragraph, level=level)
        enable_flow_control(paragraph, keep_lines=True)
        if label:
            format_run(paragraph.add_run(f"{label} "), bold=True, italic=False)
        if title:
            format_run(paragraph.add_run(f"{title} "), bold=True, italic=False)
        if text:
            format_run(paragraph.add_run(text), bold=False, italic=False)
        return paragraph
