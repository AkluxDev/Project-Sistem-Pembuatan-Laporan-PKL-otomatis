from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from .docx_utils import set_run_fonts
from .paragraph_engine import enable_flow_control
from .styles import FONT_NAME, HEADING_SIZE, apply_signature_style


class TocEngine:
    """
    TocEngine V2 — Sistem Daftar Isi Profesional.
    Menggunakan native Word fields dengan struktur XML yang kompleks
    agar terdeteksi penuh sebagai field TOC oleh Microsoft Word.
    """
    def __init__(self, document):
        self.document = document

    def apply_heading_style(self, paragraph, level: int):
        """
        Terapkan style Heading X dan pastikan outlineLvl terisi.
        """
        level = max(1, min(4, int(level)))
        style_name = f"Heading {level}"
        paragraph.style = style_name
        
        # Explicit outline level for TOC discovery
        p_pr = paragraph._p.get_or_add_pPr()
        outline_el = p_pr.find(qn("w:outlineLvl"))
        if outline_el is None:
            outline_el = OxmlElement("w:outlineLvl")
            p_pr.append(outline_el)
        outline_el.set(qn("w:val"), str(level - 1))
        
        enable_flow_control(paragraph, keep_next=True, keep_lines=True)
        return paragraph

    def insert_heading(self, text: str, level: int = 1):
        """
        Tambahkan satu paragraf heading dengan style yang benar.
        """
        paragraph = self.document.add_paragraph()
        self.apply_heading_style(paragraph, level)
        
        run = paragraph.add_run(text.strip())
        set_run_fonts(run, font_name=FONT_NAME)
        return paragraph

    def add_toc(self, title: str = "DAFTAR ISI"):
        """
        Sisipkan field TOC kompleks.
        Struktur XML: begin -> instrText -> separate -> text -> end
        """
        # 1. Judul Halaman
        title_para = self.document.add_paragraph()
        apply_signature_style(title_para)
        run_title = title_para.add_run(title.upper())
        set_run_fonts(run_title, font_name=FONT_NAME, size=HEADING_SIZE, bold=True)
        title_para.paragraph_format.space_after = Pt(24)

        # 2. Paragraf Field TOC
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run()
        
        # [BEGIN]
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        fld_begin.set(qn("w:dirty"), "true")
        
        # [INSTRUCTION]
        # \o "1-3" : Build TOC from Heading 1-3
        # \h       : Hyperlinks
        # \z       : Hide page numbers in web view
        # \u       : Use outline levels
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = r' TOC \o "1-3" \h \z \u '
        
        # [SEPARATE]
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        
        # [PLACEHOLDER]
        p_text = OxmlElement("w:t")
        p_text.text = "Klik kanan di sini lalu pilih 'Update Field' untuk menampilkan daftar isi."
        
        # [END]
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")

        # Assemble run
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_sep)
        run._r.append(p_text)
        run._r.append(fld_end)
        
        return paragraph

    def insert_toc(self, title: str = "DAFTAR ISI"):
        return self.add_toc(title)
